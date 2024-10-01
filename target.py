
import docx

# Create a new Word document
doc = docx.Document()

# Add a title
doc.add_heading('John Doe', 0)

# Add a contact section
doc.add_paragraph('123 Main Street')
doc.add_paragraph('Anytown, CA 12345')
doc.add_paragraph('(123) 456-7890')
doc.add_paragraph('john.doe@email.com')

# Add a summary section
doc.add_heading('Summary', level=1)
doc.add_paragraph('Highly motivated and experienced software engineer with a passion for building innovative solutions. Skilled in Python, Ruby, and JavaScript, with expertise in web development and data analysis.')

# Add a skills section
doc.add_heading('Skills', level=1)
doc.add_paragraph('Programming Languages: Python, Ruby, JavaScript, SQL')
doc.add_paragraph('Frameworks: Django, Rails, React')
doc.add_paragraph('Databases: PostgreSQL, MySQL')
doc.add_paragraph('Cloud Platforms: AWS, GCP')

# Add an experience section
doc.add_heading('Experience', level=1)

# Add an experience entry
doc.add_paragraph('Software Engineer', style='Heading 2')
doc.add_paragraph('Acme Corporation | Anytown, CA | 2020 - Present')
doc.add_paragraph('- Developed and maintained web applications using Python, Django, and PostgreSQL.')
doc.add_paragraph('- Implemented new features and improvements to existing systems.')
doc.add_paragraph('- Worked closely with product managers and designers to deliver high-quality software.')

# Add an education section
doc.add_heading('Education', level=1)
doc.add_paragraph('Bachelor of Science in Computer Science', style='Heading 2')
doc.add_paragraph('Stanford University | Stanford, CA | 2017')

# Save the document
doc.save('resume.docx')

# Add "Meetup with friends at 7 pm" to the todo file
with open('todo.txt', 'a') as f:
    f.write('Meetup with friends at 7 pm\n')
